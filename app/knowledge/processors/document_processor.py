# app/knowledge/processors/document_processor.py
"""
Production Document Processor for SpinScribe
Handles document upload, content extraction, and preprocessing for knowledge management.
"""

import asyncio
import logging
import mimetypes
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass
import json

# File processing imports
try:
    import PyPDF2
    import fitz  # PyMuPDF - better PDF processing
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import requests
    from bs4 import BeautifulSoup
    WEB_PROCESSING_AVAILABLE = True
except ImportError:
    WEB_PROCESSING_AVAILABLE = False

import csv
from io import StringIO, BytesIO

from app.core.config import settings
from app.database.connection import SessionLocal
from app.database.models.knowledge_item import KnowledgeItem

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    content: str
    title: str
    metadata: Dict[str, Any]
    file_hash: str
    processing_time: float
    word_count: int
    character_count: int
    language: str = "en"
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "title": self.title,
            "metadata": self.metadata,
            "file_hash": self.file_hash,
            "processing_time": self.processing_time,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "language": self.language
        }

@dataclass
class DocumentChunk:
    """Individual content chunk for processing"""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "chunk_index": self.chunk_index,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "metadata": self.metadata
        }

class DocumentProcessor:
    """
    Production-grade document processor for SpinScribe knowledge management.
    
    Supports:
    - PDF documents (with OCR fallback)
    - Word documents (.docx)
    - Text files (.txt, .md)
    - Web content (URLs, HTML)
    - Structured data (.json, .csv)
    """
    
    def __init__(self, project_id: str):
        self.project_id = project_id
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        # Processing configuration
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 100  # Overlap between chunks
        
        # Storage paths
        self.storage_path = Path(settings.storage_root_dir) / "documents" / project_id
        self.storage_path.mkdir(parents=True, exist_ok=True)
        
        # Supported formats
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.docx': self._process_docx,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.html': self._process_html,
            '.htm': self._process_html
        }
        
        self.logger.info(f"Document processor initialized for project: {project_id}")
        self._log_capabilities()
    
    def _log_capabilities(self):
        """Log available processing capabilities"""
        capabilities = []
        if PDF_AVAILABLE:
            capabilities.append("PDF processing")
        if DOCX_AVAILABLE:
            capabilities.append("DOCX processing")
        if WEB_PROCESSING_AVAILABLE:
            capabilities.append("Web content processing")
        
        self.logger.info(f"Available capabilities: {', '.join(capabilities)}")
        
        if not PDF_AVAILABLE:
            self.logger.warning("PDF processing unavailable - install PyPDF2 and PyMuPDF")
        if not DOCX_AVAILABLE:
            self.logger.warning("DOCX processing unavailable - install python-docx")
        if not WEB_PROCESSING_AVAILABLE:
            self.logger.warning("Web processing unavailable - install requests and beautifulsoup4")
    
    async def process_file(self, 
                          file_path: Union[str, Path],
                          knowledge_type: str = "document",
                          metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process a file and extract content
        
        Args:
            file_path: Path to the file to process
            knowledge_type: Type of knowledge item
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument with extracted content and metadata
        """
        start_time = datetime.utcnow()
        
        try:
            file_path = Path(file_path)
            
            # Validate file
            self._validate_file(file_path)
            
            # Determine file type and processor
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Read file content
            file_content = self._read_file(file_path)
            
            # Generate file hash
            file_hash = self._generate_file_hash(file_content)
            
            # Process content based on file type
            processor = self.supported_formats[file_extension]
            extracted_content, doc_metadata = await processor(file_content, file_path)
            
            # Clean and normalize content
            cleaned_content = self._clean_content(extracted_content)
            
            # Generate title
            title = self._generate_title(file_path, doc_metadata, cleaned_content)
            
            # Combine metadata
            combined_metadata = {
                "file_name": file_path.name,
                "file_size": len(file_content),
                "file_type": file_extension,
                "knowledge_type": knowledge_type,
                "processing_timestamp": datetime.utcnow().isoformat(),
                **doc_metadata,
                **(metadata or {})
            }
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Create processed document
            processed_doc = ProcessedDocument(
                content=cleaned_content,
                title=title,
                metadata=combined_metadata,
                file_hash=file_hash,
                processing_time=processing_time,
                word_count=len(cleaned_content.split()),
                character_count=len(cleaned_content)
            )
            
            self.logger.info(f"Successfully processed {file_path.name} ({processing_time:.2f}s)")
            
            return processed_doc
            
        except Exception as e:
            self.logger.error(f"Failed to process {file_path}: {e}")
            raise
    
    async def process_url(self,
                         url: str,
                         knowledge_type: str = "web_content",
                         metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process content from a URL
        
        Args:
            url: URL to process
            knowledge_type: Type of knowledge item
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument with extracted content
        """
        if not WEB_PROCESSING_AVAILABLE:
            raise ImportError("Web processing requires 'requests' and 'beautifulsoup4'")
        
        start_time = datetime.utcnow()
        
        try:
            self.logger.info(f"Processing URL: {url}")
            
            # Fetch content
            headers = {
                'User-Agent': 'SpinScribe Document Processor 1.0'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Generate hash from content
            file_hash = hashlib.md5(response.content).hexdigest()
            
            # Process HTML content
            extracted_content, doc_metadata = await self._process_html(response.content, url)
            
            # Clean content
            cleaned_content = self._clean_content(extracted_content)
            
            # Generate title
            title = doc_metadata.get("title", self._extract_domain_from_url(url))
            
            # Combine metadata
            combined_metadata = {
                "source_url": url,
                "content_type": response.headers.get("content-type", "text/html"),
                "response_status": response.status_code,
                "content_length": len(response.content),
                "knowledge_type": knowledge_type,
                "processing_timestamp": datetime.utcnow().isoformat(),
                **doc_metadata,
                **(metadata or {})
            }
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Create processed document
            processed_doc = ProcessedDocument(
                content=cleaned_content,
                title=title,
                metadata=combined_metadata,
                file_hash=file_hash,
                processing_time=processing_time,
                word_count=len(cleaned_content.split()),
                character_count=len(cleaned_content)
            )
            
            self.logger.info(f"Successfully processed URL {url} ({processing_time:.2f}s)")
            
            return processed_doc
            
        except Exception as e:
            self.logger.error(f"Failed to process URL {url}: {e}")
            raise
    
    def chunk_document(self, 
                      processed_doc: ProcessedDocument,
                      chunk_size: Optional[int] = None,
                      overlap: Optional[int] = None) -> List[DocumentChunk]:
        """
        Split document into chunks for processing
        
        Args:
            processed_doc: Processed document to chunk
            chunk_size: Size of each chunk (characters)
            overlap: Overlap between chunks
            
        Returns:
            List of document chunks
        """
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap
        
        content = processed_doc.content
        chunks = []
        
        # Handle short documents
        if len(content) <= chunk_size:
            chunk = DocumentChunk(
                content=content,
                chunk_index=0,
                start_char=0,
                end_char=len(content),
                metadata={
                    **processed_doc.metadata,
                    "chunk_info": {
                        "total_chunks": 1,
                        "chunk_method": "single"
                    }
                }
            )
            return [chunk]
        
        # Split by sentences first to avoid breaking mid-sentence
        sentences = self._split_into_sentences(content)
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Create chunk
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                    metadata={
                        **processed_doc.metadata,
                        "chunk_info": {
                            "chunk_method": "sentence_based",
                            "chunk_size": len(current_chunk)
                        }
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-overlap:] if overlap > 0 else ""
                current_chunk = overlap_text + sentence
                current_start = chunks[-1].end_char - len(overlap_text)
                chunk_index += 1
            else:
                current_chunk += sentence
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
                metadata={
                    **processed_doc.metadata,
                    "chunk_info": {
                        "chunk_method": "sentence_based",
                        "chunk_size": len(current_chunk),
                        "is_final_chunk": True
                    }
                }
            )
            chunks.append(chunk)
        
        # Update total chunks info
        for chunk in chunks:
            chunk.metadata["chunk_info"]["total_chunks"] = len(chunks)
        
        self.logger.info(f"Split document into {len(chunks)} chunks")
        return chunks
    
    # File type processors
    
    async def _process_pdf(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process PDF files"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing requires PyPDF2 and PyMuPDF")
        
        try:
            # Try PyMuPDF first (better text extraction)
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            
            metadata = {
                "pages": len(doc),
                "extraction_method": "pymupdf"
            }
            
        except Exception as e:
            self.logger.warning(f"PyMuPDF failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
                
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "extraction_method": "pypdf2"
                }
                
            except Exception as e2:
                raise ValueError(f"PDF processing failed with both methods: {e}, {e2}")
        
        return text_content, metadata
    
    async def _process_text(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process plain text files"""
        try:
            # Try UTF-8 first
            content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file with any supported encoding")
        
        metadata = {
            "encoding": "utf-8",
            "file_type": "text/plain"
        }
        
        return content, metadata
    
    async def _process_markdown(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown files"""
        content = file_content.decode('utf-8')
        
        # Extract title from first heading if available
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else None
        
        metadata = {
            "file_type": "text/markdown",
            "has_title": bool(title),
            "title": title
        }
        
        return content, metadata
    
    async def _process_docx(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing requires python-docx")
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            table_content = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_content += " | ".join(row_text) + "\n"
            
            if table_content:
                text_content += "\n\nTable Data:\n" + table_content
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            
            return text_content, metadata
            
        except Exception as e:
            raise ValueError(f"DOCX processing failed: {e}")
    
    async def _process_json(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process JSON files"""
        try:
            json_data = json.loads(file_content.decode('utf-8'))
            
            # Convert JSON to readable text
            if isinstance(json_data, dict):
                text_content = self._dict_to_text(json_data)
            elif isinstance(json_data, list):
                text_content = self._list_to_text(json_data)
            else:
                text_content = str(json_data)
            
            metadata = {
                "file_type": "application/json",
                "json_type": type(json_data).__name__,
                "data_structure": self._analyze_json_structure(json_data)
            }
            
            return text_content, metadata
            
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file: {e}")
    
    async def _process_csv(self, file_content: bytes, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process CSV files"""
        try:
            content = file_content.decode('utf-8')
            csv_reader = csv.reader(StringIO(content))
            
            rows = list(csv_reader)
            if not rows:
                return "", {"file_type": "text/csv", "rows": 0, "columns": 0}
            
            # Convert to readable text
            headers = rows[0] if rows else []
            text_content = f"CSV Data with columns: {', '.join(headers)}\n\n"
            
            for i, row in enumerate(rows[:100]):  # Limit to first 100 rows
                text_content += f"Row {i+1}: {' | '.join(row)}\n"
            
            if len(rows) > 100:
                text_content += f"\n... and {len(rows) - 100} more rows"
            
            metadata = {
                "file_type": "text/csv",
                "rows": len(rows),
                "columns": len(headers),
                "headers": headers
            }
            
            return text_content, metadata
            
        except Exception as e:
            raise ValueError(f"CSV processing failed: {e}")
    
    async def _process_html(self, file_content: bytes, source: Union[Path, str]) -> Tuple[str, Dict[str, Any]]:
        """Process HTML content"""
        if not WEB_PROCESSING_AVAILABLE:
            raise ImportError("HTML processing requires beautifulsoup4")
        
        try:
            soup = BeautifulSoup(file_content, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else None
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract main content
            # Try to find main content areas
            main_content = soup.find('main') or soup.find('article') or soup.find('body')
            
            if main_content:
                text_content = main_content.get_text()
            else:
                text_content = soup.get_text()
            
            # Clean up whitespace
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            # Extract metadata
            meta_description = ""
            meta_keywords = ""
            
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                meta_description = meta_desc.get('content', '')
            
            meta_kw = soup.find('meta', attrs={'name': 'keywords'})
            if meta_kw:
                meta_keywords = meta_kw.get('content', '')
            
            metadata = {
                "file_type": "text/html",
                "title": title,
                "meta_description": meta_description,
                "meta_keywords": meta_keywords,
                "source": str(source)
            }
            
            return text_content, metadata
            
        except Exception as e:
            raise ValueError(f"HTML processing failed: {e}")
    
    # Utility methods
    
    def _validate_file(self, file_path: Path):
        """Validate file for processing"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        if file_size == 0:
            raise ValueError("File is empty")
    
    def _read_file(self, file_path: Path) -> bytes:
        """Read file content"""
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except Exception as e:
            raise IOError(f"Failed to read file: {e}")
    
    def _generate_file_hash(self, content: bytes) -> str:
        """Generate MD5 hash of file content"""
        return hashlib.md5(content).hexdigest()
    
    def _clean_content(self, content: str) -> str:
        """Clean and normalize extracted content"""
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove excessive newlines
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Trim whitespace
        content = content.strip()
        
        return content
    
    def _generate_title(self, file_path: Path, metadata: Dict[str, Any], content: str) -> str:
        """Generate appropriate title for document"""
        # Try metadata title first
        if metadata.get("title"):
            return metadata["title"]
        
        # Try first line of content
        first_line = content.split('\n')[0].strip()
        if first_line and len(first_line) < 100:
            return first_line
        
        # Fall back to filename
        return file_path.stem.replace('_', ' ').replace('-', ' ').title()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting (can be enhanced with NLTK)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() + ' ' for s in sentences if s.strip()]
    
    def _dict_to_text(self, data: dict, indent: int = 0) -> str:
        """Convert dictionary to readable text"""
        text = ""
        for key, value in data.items():
            prefix = "  " * indent
            if isinstance(value, dict):
                text += f"{prefix}{key}:\n{self._dict_to_text(value, indent + 1)}"
            elif isinstance(value, list):
                text += f"{prefix}{key}: {self._list_to_text(value, indent + 1)}"
            else:
                text += f"{prefix}{key}: {value}\n"
        return text
    
    def _list_to_text(self, data: list, indent: int = 0) -> str:
        """Convert list to readable text"""
        text = ""
        prefix = "  " * indent
        for i, item in enumerate(data[:50]):  # Limit to 50 items
            if isinstance(item, dict):
                text += f"{prefix}Item {i+1}:\n{self._dict_to_text(item, indent + 1)}"
            else:
                text += f"{prefix}- {item}\n"
        
        if len(data) > 50:
            text += f"{prefix}... and {len(data) - 50} more items\n"
        
        return text
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure for metadata"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys())[:20],  # First 20 keys
                "total_keys": len(data)
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_types": list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {
                "type": type(data).__name__,
                "value": str(data)[:100]
            }
    
    def _extract_domain_from_url(self, url: str) -> str:
        """Extract domain name from URL for title"""
        import re
        match = re.search(r'https?://([^/]+)', url)
        return match.group(1) if match else url

# Factory function
def create_document_processor(project_id: str) -> DocumentProcessor:
    """
    Factory function to create a DocumentProcessor instance
    
    Args:
        project_id: Project ID for processor
        
    Returns:
        Initialized DocumentProcessor
    """
    return DocumentProcessor(project_id)

# Export main classes
__all__ = [
    'DocumentProcessor',
    'ProcessedDocument', 
    'DocumentChunk',
    'create_document_processor'
]